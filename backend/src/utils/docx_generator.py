"""DOCX Resume Generator - Classic Times New Roman Template"""

from typing import List

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import io


def add_horizontal_line(paragraph: Paragraph) -> None:
    """Add a horizontal line (border) below a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # Line thickness
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_font(
    run: Run,
    font_name: str = "Times New Roman",
    size: int = 11,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_header_row(
    table: Table,
    left_text: str,
    right_text: str,
    *,
    bold: bool = True,
    italic: bool = False,
) -> None:
    """Add a row with left-aligned and right-aligned text"""
    row = table.add_row()
    left_cell = row.cells[0]
    right_cell = row.cells[1]

    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run(left_text)
    set_font(left_run, bold=bold, italic=italic)

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run(right_text)
    set_font(right_run, bold=bold, italic=italic)


def html_to_docx(html_content: str) -> bytes:
    """
    Convert HTML resume to DOCX format.

    Args:
        html_content: HTML string of the resume

    Returns:
        DOCX file as bytes
    """
    soup = BeautifulSoup(html_content, "html.parser")
    doc: DocxDocument = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Extract header info
    header_div = soup.find("div", class_="header")
    if header_div:
        # Name
        h1 = header_div.find("h1")
        if h1:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(h1.get_text(strip=True))
            set_font(name_run, size=16, bold=True)

        # Contact info
        contact_div = header_div.find("div", class_="contact-info")
        if contact_div:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(contact_div.get_text(strip=True))
            set_font(contact_run, size=11)

    # Process sections
    for h2 in soup.find_all("h2"):
        section_title = h2.get_text(strip=True).upper()

        # Add section header
        section_header = doc.add_paragraph()
        section_run = section_header.add_run(section_title)
        set_font(section_run, size=11, bold=True)
        add_horizontal_line(section_header)
        section_header.paragraph_format.space_after = Pt(4)
        section_header.paragraph_format.space_before = Pt(8)

        # Find all items in this section
        current = h2.find_next_sibling()

        while current is not None:
            if isinstance(current, Tag):
                if current.name == "h2":
                    break

                class_attr = current.get("class")
                class_list: List[str] = (
                    class_attr if isinstance(class_attr, list) else []
                )

                if "education-item" in class_list:
                    process_item(doc, current, "education")
                elif "experience-item" in class_list:
                    process_item(doc, current, "experience")
                elif "project-item" in class_list:
                    process_item(doc, current, "project")
                elif "org-item" in class_list:
                    process_item(doc, current, "org")
                elif "award-item" in class_list:
                    process_item(doc, current, "award")
                elif "skills-section" in class_list:
                    for skill_line in current.find_all("div", class_="skill-line"):
                        para = doc.add_paragraph()
                        strong = skill_line.find("strong")
                        if strong:
                            label_text = strong.get_text(strip=True)
                            label_run = para.add_run(label_text + " ")
                            set_font(label_run, size=11, bold=True)
                            # Get text after strong tag
                            remaining_text = (
                                skill_line.get_text(strip=True)
                                .replace(label_text, "", 1)
                                .lstrip(": ")
                            )
                            content_run = para.add_run(remaining_text)
                            set_font(content_run, size=11)

            current = current.find_next_sibling()

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def process_item(doc: DocxDocument, item_div: Tag, item_type: str) -> None:
    """Process an education/experience/project/org/award item"""
    # Header row (title and location)
    item_header = item_div.find("div", class_="item-header")
    if item_header:
        title_span = item_header.find("span", class_="item-title")
        location_span = item_header.find("span", class_="item-location")

        if title_span and location_span:
            table = doc.add_table(rows=0, cols=2)
            add_header_row(
                table,
                title_span.get_text(strip=True),
                location_span.get_text(strip=True),
                bold=True,
            )

    # Description (for org items)
    description = item_div.find("div", class_="description")
    if description:
        desc_para = doc.add_paragraph(description.get_text(strip=True))
        if desc_para.runs:
            set_font(desc_para.runs[0], size=11)

    # Subtitle (for projects - technologies)
    subtitle = item_div.find("div", class_=["item-subtitle", "project-subtitle"])
    if subtitle:
        # This goes in dates row for projects, or separate for experience
        pass  # Handle below with dates

    # Dates row
    item_dates = item_div.find("div", class_="item-dates")
    if item_dates:
        dates_left = item_dates.find("span", class_="item-dates-left")
        dates_right = item_dates.find("span", class_="item-dates-right")

        if dates_left and dates_right:
            table = doc.add_table(rows=0, cols=2)
            add_header_row(
                table,
                dates_left.get_text(strip=True),
                dates_right.get_text(strip=True),
                bold=False,
                italic=True,
            )
    elif subtitle:
        # For projects without separate dates row
        subtitle_para = doc.add_paragraph(subtitle.get_text(strip=True))
        if subtitle_para.runs:
            subtitle_para.runs[0].italic = True
            set_font(subtitle_para.runs[0], size=11)

    # Bullets
    ul = item_div.find("ul")
    if ul:
        for li in ul.find_all("li"):
            bullet_para = doc.add_paragraph(
                li.get_text(strip=True), style="List Bullet"
            )
            if bullet_para.runs:
                set_font(bullet_para.runs[0], size=11)

"""Page estimation and control utilities for DOCX resume generation."""

from typing import List, Tuple, Dict, Any
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class PageEstimator:
    """Estimates and controls page count for DOCX documents."""
    
    # Standard page dimensions (Letter size with 0.75" margins)
    PAGE_WIDTH_INCHES = 8.5  # Letter width
    PAGE_HEIGHT_INCHES = 11.0  # Letter height
    MARGIN_INCHES = 0.75
    
    # Effective content area
    CONTENT_WIDTH_INCHES = PAGE_WIDTH_INCHES - (2 * MARGIN_INCHES)  # 7.0"
    CONTENT_HEIGHT_INCHES = PAGE_HEIGHT_INCHES - (2 * MARGIN_INCHES)  # 9.5"
    
    # Character counts per line with Times New Roman 11pt
    CHARS_PER_LINE = 75  # Reduced from 85 - more realistic for 7" width at 11pt
    LINES_PER_PAGE = 35  # Reduced from 55 - more realistic for 9.5" height with spacing
    
    # Spacing multipliers for different elements (increased for accuracy)
    SECTION_HEADER_SPACING = 3.0  # Lines consumed by section header + line + spacing
    ITEM_HEADER_SPACING = 2.0  # Lines for item title/location + spacing
    ITEM_SUBTITLE_SPACING = 1.5  # Lines for dates/role + spacing
    BULLET_SPACING = 1.5  # Lines per bullet point + spacing
    CONTACT_SPACING = 3.0  # Lines for contact info + spacing
    
    @classmethod
    def estimate_text_lines(cls, text: str) -> float:
        """Estimate number of lines a text block will occupy."""
        if not text:
            return 0
        
        # Remove extra whitespace and split by actual newlines
        clean_text = re.sub(r'\s+', ' ', text.strip())
        lines = clean_text.split('\n')
        
        total_lines = 0
        for line in lines:
            if len(line) <= cls.CHARS_PER_LINE:
                total_lines += 1
            else:
                # Estimate line wraps for long text (more conservative)
                wrapped_lines = max(1, (len(line) + cls.CHARS_PER_LINE - 1) // cls.CHARS_PER_LINE)
                total_lines += wrapped_lines
        
        return total_lines * 1.2  # Add 20% buffer for formatting overhead
    
    @classmethod
    def estimate_section_lines(cls, section_data: Dict[str, Any]) -> float:
        """Estimate lines consumed by a resume section."""
        lines = 0
        
        # Section header
        lines += cls.SECTION_HEADER_SPACING
        
        # Process items in section
        items = section_data.get('items', [])
        for item in items:
            # Item header (title + location)
            title = item.get('title', '')
            location = item.get('location', '')
            lines += cls.ITEM_HEADER_SPACING
            
            # Subtitle/dates
            subtitle = item.get('subtitle', '')
            dates = item.get('dates', '')
            if subtitle or dates:
                lines += cls.ITEM_SUBTITLE_SPACING
            
            # Bullet points
            bullets = item.get('bullets', [])
            for bullet in bullets:
                lines += cls.estimate_text_lines(bullet) * cls.BULLET_SPACING
        
        return lines
    
    @classmethod
    def estimate_resume_pages(cls, resume_data: Dict[str, Any]) -> float:
        """Estimate total pages for a complete resume."""
        total_lines = 0
        
        # Contact info (name + contact)
        total_lines += cls.CONTACT_SPACING
        
        # Process each section
        sections = resume_data.get('sections', [])
        for section in sections:
            total_lines += cls.estimate_section_lines(section)
        
        # Convert to pages with calibration factor based on real data
        estimated_pages = (total_lines / cls.LINES_PER_PAGE) * 1.0  # Updated calibration: 1.3 * 0.79 ≈ 1.0
        return estimated_pages
    
    @classmethod
    def calibrate_estimation(cls, estimated_pages: float, actual_pages: float) -> float:
        """Calculate calibration factor based on actual vs estimated pages."""
        if estimated_pages <= 0:
            return 1.0
        return actual_pages / estimated_pages
    
    @classmethod
    def calculate_spacing_adjustment(cls, current_pages: float, target_pages: float = 2.0) -> Dict[str, float]:
        """Calculate spacing adjustments to fit target page count."""
        if abs(current_pages - target_pages) < 0.1:  # Within 10% of target
            return {'adjustment_factor': 1.0, 'space_before': Pt(6), 'space_after': Pt(6)}
        
        # Calculate needed adjustment
        page_diff = current_pages - target_pages
        adjustment_factor = target_pages / current_pages
        
        # Generate spacing recommendations
        if page_diff > 0:  # Too many pages, reduce spacing
            space_before = max(Pt(2), Pt(6 * adjustment_factor))
            space_after = max(Pt(2), Pt(6 * adjustment_factor))
            bullet_spacing = max(Pt(0), Pt(3 * adjustment_factor))
        else:  # Too few pages, increase spacing
            space_before = min(Pt(12), Pt(6 * adjustment_factor))
            space_after = min(Pt(12), Pt(6 * adjustment_factor))
            bullet_spacing = min(Pt(8), Pt(3 * adjustment_factor))
        
        return {
            'adjustment_factor': adjustment_factor,
            'space_before': space_before,
            'space_after': space_after,
            'bullet_spacing': bullet_spacing,
            'section_spacing_before': min(Pt(10), Pt(8 * adjustment_factor)),
            'section_spacing_after': min(Pt(8), Pt(4 * adjustment_factor))
        }


def add_horizontal_line(paragraph):
    """Add a horizontal line (border) below a paragraph."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_font(run, font_name='Times New Roman', size=11, bold=False, italic=False):
    """Set font properties for a run."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_header_row(table, left_text, right_text, bold=True, italic=False):
    """Add a row with left-aligned and right-aligned text."""
    row = table.add_row()
    left_cell, right_cell = row.cells[0], row.cells[1]
    left_run = left_cell.paragraphs[0].add_run(left_text)
    set_font(left_run, bold=bold, italic=italic)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run(right_text)
    set_font(right_run, bold=bold, italic=italic)


def apply_spacing_adjustments(paragraph, spacing_config: Dict[str, Any]):
    """Apply calculated spacing adjustments to a paragraph."""
    if 'space_before' in spacing_config:
        paragraph.paragraph_format.space_before = spacing_config['space_before']
    if 'space_after' in spacing_config:
        paragraph.paragraph_format.space_after = spacing_config['space_after']


def create_section_header(doc, text, spacing_config: Dict[str, Any]):
    """Create a section header with proper spacing."""
    section_header = doc.add_paragraph()
    set_font(section_header.add_run(text.upper()), size=11, bold=True)
    add_horizontal_line(section_header)
    apply_spacing_adjustments(section_header, {
        'space_before': spacing_config.get('section_spacing_before', Pt(8)),
        'space_after': spacing_config.get('section_spacing_after', Pt(4))
    })
    return section_header


def add_bullet_point(doc, text, spacing_config: Dict[str, Any]):
    """Add a bullet point with adjusted spacing."""
    bullet = doc.add_paragraph(text, style='List Bullet')
    set_font(bullet.runs[0], size=11)
    apply_spacing_adjustments(bullet, {
        'space_before': Pt(0),
        'space_after': spacing_config.get('bullet_spacing', Pt(3))
    })
    return bullet


def get_default_spacing_config() -> Dict[str, Any]:
    """Return a default spacing configuration for 2-3 page resumes.

    This is a convenience for agents so they can use a single dict
    instead of hard-coding Pt() values everywhere.
    """
    return {
        'space_before': Pt(6),
        'space_after': Pt(6),
        'bullet_spacing': Pt(3),
        'section_spacing_before': Pt(8),
        'section_spacing_after': Pt(4),
    }


def get_compact_spacing_config() -> Dict[str, Any]:
    """Return a spacing configuration for slightly more compact layouts."""
    return {
        'space_before': Pt(4),
        'space_after': Pt(4),
        'bullet_spacing': Pt(2),
        'section_spacing_before': Pt(6),
        'section_spacing_after': Pt(3),
    }


def get_spacious_spacing_config() -> Dict[str, Any]:
    """Return a spacing configuration for slightly more spacious layouts."""
    return {
        'space_before': Pt(8),
        'space_after': Pt(8),
        'bullet_spacing': Pt(4),
        'section_spacing_before': Pt(10),
        'section_spacing_after': Pt(6),
    }


def add_experience_block(
    doc: Document,
    *,
    company: str,
    location: str,
    role: str,
    dates: str,
    bullets: list[str],
    spacing_config: Dict[str, Any],
) -> None:
    """Add a standard experience block (header + bullets).

    This helper keeps layout consistent while allowing the agent to focus
    on content. It uses tables for aligned headers and the shared spacing
    configuration for vertical rhythm.
    """
    table = doc.add_table(rows=0, cols=2)
    add_header_row(table, company, location, bold=True)
    add_header_row(table, role, dates, bold=False, italic=True)

    # Spacing paragraph after the header rows
    spacing_para = doc.add_paragraph()
    apply_spacing_adjustments(spacing_para, {
        'space_before': spacing_config.get('space_before', Pt(6)),
        'space_after': spacing_config.get('space_after', Pt(6)),
    })

    for bullet_text in bullets:
        add_bullet_point(doc, bullet_text, spacing_config)


def add_education_block(
    doc: Document,
    *,
    institution: str,
    location: str,
    degree: str,
    dates: str,
    bullets: list[str],
    spacing_config: Dict[str, Any],
) -> None:
    """Add a standard education block (institution + degree + bullets)."""
    table = doc.add_table(rows=0, cols=2)
    add_header_row(table, institution, location, bold=True)
    add_header_row(table, degree, dates, bold=False, italic=True)

    spacing_para = doc.add_paragraph()
    apply_spacing_adjustments(spacing_para, {
        'space_before': spacing_config.get('space_before', Pt(6)),
        'space_after': spacing_config.get('space_after', Pt(6)),
    })

    for bullet_text in bullets:
        add_bullet_point(doc, bullet_text, spacing_config)

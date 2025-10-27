"""Parse marked-up plain text and convert to DOCX"""

from typing import Dict, List, Optional, Union

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re


ResumeItem = Dict[str, Union[str, List[str], None]]


def add_horizontal_line(paragraph):
    """Add a horizontal line (border) below a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_font(run, font_name="Times New Roman", size=11, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_header_row(table, left_text, right_text, bold=True, italic=False):
    """Add a row with left-aligned and right-aligned text"""
    row = table.add_row()
    left_cell = row.cells[0]
    right_cell = row.cells[1]

    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run(left_text)
    set_font(left_run, bold=bold, italic=italic)

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run(right_text)
    set_font(right_run, bold=bold, italic=italic)

    return row


def parse_marked_text_to_docx(marked_text: str) -> bytes:
    """
    Parse marked-up plain text and convert to DOCX.

    Args:
        marked_text: Plain text with section markers

    Returns:
        DOCX file as bytes
    """
    doc: DocxDocument = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = marked_text.strip().split("\n")
    current_section: Optional[str] = None
    current_item: ResumeItem = {}

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # NAME
        if line == "[NAME]":
            i += 1
            if i < len(lines):
                name_para = doc.add_paragraph()
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                name_run = name_para.add_run(lines[i].strip())
                set_font(name_run, size=16, bold=True)
            i += 1
            continue

        # CONTACT
        if line == "[CONTACT]":
            i += 1
            if i < len(lines):
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_run = contact_para.add_run(lines[i].strip())
                set_font(contact_run, size=11)
            i += 1
            continue

        # SECTION
        if line == "[SECTION]":
            # Finalize previous item if any
            if current_item:
                _process_item(doc, current_item)
                current_item = {}

            i += 1
            if i < len(lines):
                current_section = lines[i].strip()
                section_header = doc.add_paragraph()
                section_run = section_header.add_run(current_section)
                set_font(section_run, size=11, bold=True)
                add_horizontal_line(section_header)
                section_header.paragraph_format.space_after = Pt(4)
                section_header.paragraph_format.space_before = Pt(8)
            i += 1
            continue

        # ITEM
        if line == "[ITEM]":
            # Finalize previous item if any
            if current_item:
                _process_item(doc, current_item)
            current_item = {
                "section": current_section,
                "bullets": [],
                "skill_lines": [],
            }
            i += 1
            continue

        # Parse item fields
        if line.startswith("[HEADER_LEFT]"):
            current_item["header_left"] = line.replace("[HEADER_LEFT]", "").strip()
        elif line.startswith("[HEADER_RIGHT]"):
            current_item["header_right"] = line.replace("[HEADER_RIGHT]", "").strip()
        elif line.startswith("[DATES_LEFT]"):
            current_item["dates_left"] = line.replace("[DATES_LEFT]", "").strip()
        elif line.startswith("[DATES_RIGHT]"):
            current_item["dates_right"] = line.replace("[DATES_RIGHT]", "").strip()
        elif line.startswith("[SUBTITLE]"):
            current_item["subtitle"] = line.replace("[SUBTITLE]", "").strip()
        elif line.startswith("[DESCRIPTION]"):
            current_item["description"] = line.replace("[DESCRIPTION]", "").strip()
        elif line.startswith("[BULLET]"):
            bullets_value = current_item.get("bullets")
            if isinstance(bullets_value, list):
                bullets_list: List[str] = bullets_value
            else:
                bullets_list = []
                current_item["bullets"] = bullets_list
            bullets_list.append(line.replace("[BULLET]", "").strip())
        elif line.startswith("[SKILL_LINE]"):
            skills_value = current_item.get("skill_lines")
            if isinstance(skills_value, list):
                skill_lines_list: List[str] = skills_value
            else:
                skill_lines_list = []
                current_item["skill_lines"] = skill_lines_list
            skill_lines_list.append(line.replace("[SKILL_LINE]", "").strip())

        i += 1

    # Finalize last item
    if current_item:
        _process_item(doc, current_item)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def _process_item(doc: DocxDocument, item: ResumeItem) -> None:
    """Process a single item and add to document"""
    section_value = item.get("section")
    section = section_value if isinstance(section_value, str) else ""

    # Skills section special handling
    if section == "SKILLS & INTERESTS":
        skill_lines = item.get("skill_lines")
        if isinstance(skill_lines, list):
            for skill_line in skill_lines:
                para = doc.add_paragraph()
                # Parse "Skills: ..." format
                if ":" in skill_line:
                    label, content = skill_line.split(":", 1)
                    label_run = para.add_run(label.strip() + ": ")
                    set_font(label_run, size=11, bold=True)
                    content_run = para.add_run(content.strip())
                    set_font(content_run, size=11)
                else:
                    run = para.add_run(skill_line)
                    set_font(run, size=11)
        return

    # Header row (title and location)
    header_left = item.get("header_left")
    header_right = item.get("header_right")
    if isinstance(header_left, str) and isinstance(header_right, str):
        table = doc.add_table(rows=0, cols=2)
        add_header_row(table, header_left, header_right, bold=True)

    # Description (for org items)
    description = item.get("description")
    if isinstance(description, str):
        desc_para = doc.add_paragraph(description)
        if desc_para.runs:
            set_font(desc_para.runs[0], size=11)

    # Dates row
    dates_left = item.get("dates_left")
    dates_right = item.get("dates_right")
    if isinstance(dates_left, str) and isinstance(dates_right, str):
        table = doc.add_table(rows=0, cols=2)
        add_header_row(table, dates_left, dates_right, bold=False, italic=True)
    # Subtitle (for projects without dates row)
    else:
        subtitle = item.get("subtitle")
        if isinstance(subtitle, str):
            subtitle_para = doc.add_paragraph(subtitle)
            if subtitle_para.runs:
                subtitle_para.runs[0].italic = True
                set_font(subtitle_para.runs[0], size=11)

    # Bullets
    bullets = item.get("bullets")
    if isinstance(bullets, list):
        for bullet in bullets:
            bullet_para = doc.add_paragraph(bullet, style="List Bullet")
            if bullet_para.runs:
                set_font(bullet_para.runs[0], size=11)

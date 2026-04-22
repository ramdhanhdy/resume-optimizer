"""Deterministic DOCX renderer for plain-text resume previews."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


COMMON_SECTION_HEADINGS = {
    "summary",
    "professional summary",
    "profile",
    "experience",
    "work experience",
    "professional experience",
    "employment history",
    "education",
    "skills",
    "technical skills",
    "projects",
    "certifications",
    "awards",
    "publications",
    "leadership",
    "activities",
    "languages",
    "interests",
    "volunteer",
    "volunteer experience",
}


def plain_text_to_docx(resume_text: str) -> bytes:
    """Render a plain-text resume into a conservative DOCX document."""
    document: DocxDocument = Document()
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = [line.rstrip() for line in resume_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    non_empty = [line for line in lines if line.strip()]
    if not non_empty:
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return output.read()

    name = non_empty[0].strip()
    name_para = document.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    _set_font(name_run, size=16, bold=True)

    idx = 1
    while idx < len(non_empty) and not _is_section_heading(non_empty[idx].strip()):
        contact_para = document.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(non_empty[idx].strip())
        _set_font(contact_run, size=11)
        idx += 1

    while idx < len(non_empty):
        line = non_empty[idx].strip()
        if _is_section_heading(line):
            header = document.add_paragraph()
            header_run = header.add_run(line)
            _set_font(header_run, size=11, bold=True)
            _add_horizontal_line(header)
            header.paragraph_format.space_before = Pt(8)
            header.paragraph_format.space_after = Pt(4)
        elif _is_bullet_line(line):
            para = document.add_paragraph(_strip_bullet_prefix(line), style="List Bullet")
            if para.runs:
                _set_font(para.runs[0], size=11)
        else:
            para = document.add_paragraph(line)
            if para.runs:
                _set_font(para.runs[0], size=11)
            para.paragraph_format.space_after = Pt(2)
        idx += 1

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()


def _add_horizontal_line(paragraph) -> None:
    p = paragraph._element
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_font(run, *, size: int = 11, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold


def _is_section_heading(line: str) -> bool:
    line = line.strip()
    if not line:
        return False
    if len(line) > 60:
        return False
    if "@" in line or "http" in line.lower():
        return False
    alpha = [ch for ch in line if ch.isalpha()]
    if not alpha:
        return False

    normalized = re.sub(r"\s+", " ", re.sub(r"[^A-Za-z]+", " ", line)).strip().lower()
    if normalized in COMMON_SECTION_HEADINGS:
        return True

    if line == line.upper():
        return True

    return bool(re.fullmatch(r"[A-Z][a-z]+", line))


def _is_bullet_line(line: str) -> bool:
    return line.startswith(("•", "-", "*"))


def _strip_bullet_prefix(line: str) -> str:
    return line[1:].strip() if _is_bullet_line(line) else line

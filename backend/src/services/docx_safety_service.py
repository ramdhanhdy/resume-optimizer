"""Safety checks for LLM-generated DOCX Python templates using gpt-oss-safeguard-20b.

This module integrates the OpenAI gpt-oss-safeguard-20b safety reasoning model
via OpenRouter to classify generated Python code before it is executed in the
DOCX export sandbox.

The safeguard is **opt-in** and controlled by the DOCX_SAFEGUARD_ENABLED
environment variable. When disabled or misconfigured, the existing AST-based
sandbox remains the only line of defense.
"""

from __future__ import annotations

import asyncio
import json
import os
from typing import Any, Dict, Optional

from openai import OpenAI


DOCX_SAFEGUARD_ENABLED = os.getenv("DOCX_SAFEGUARD_ENABLED", "false").lower() == "true"
DOCX_SAFEGUARD_MODEL = os.getenv("DOCX_SAFEGUARD_MODEL", "openai/gpt-oss-safeguard-20b")


# Policy tailored for DOCX template Python code safety in this app.
_DOCX_CODE_POLICY = """# Policy: DOCX Template Python Code Safety for Resume Optimizer

## INSTRUCTIONS

You are gpt-oss-safeguard-20b, a safety reasoning model. Your task is to
classify a single Python code snippet that generates a DOCX resume template.

Apply ONLY the rules in this policy. Do not make up new rules.

For each snippet, decide whether it is:
- SAFE
- UNSAFE
- REVIEW (only when you cannot confidently decide SAFE vs UNSAFE)

You MUST respond with a single JSON object with this exact shape:

{
  "label": "SAFE" | "UNSAFE" | "REVIEW",
  "reasoning": "<short natural language explanation>",
  "violations": [
    "<short description of each relevant policy violation or risk>"
  ]
}

- `reasoning` should be concise but specific.
- `violations` can be an empty array for SAFE code.
- Never include the raw code in `reasoning` or `violations`.
- Do NOT output anything that is not valid JSON.
- Use medium-depth reasoning effort: be careful but not excessively verbose.

## DEFINITIONS

- DOCX template code: Python code whose purpose is to build a resume inside a
  `python-docx` `Document` object and then save it.
- Sandbox: A restricted Python execution environment with limited builtins and
  modules.
- OS-level operations: Any operations that interact with the underlying
  operating system, filesystem, environment variables, or processes.
- Network operations: Any operations that send or receive data over a network
  (HTTP, HTTPS, sockets, RPC, databases, queues).

## VIOLATES (UNSAFE)

Label the code as UNSAFE if it clearly attempts any of the following:

1. Access the operating system, filesystem, or environment:
   - Import or use modules like: `os`, `os.path`, `subprocess`, `pathlib`,
     `sys`, `shutil`, `socket`, `threading`, `multiprocessing`, `ctypes`,
     `inspect`.
   - Read, write, or delete files (e.g., `open(...)`, `Path(...).read_text()`,
     `os.remove`, `os.listdir`, `os.walk`, etc.).
   - Access environment variables (e.g., `os.environ`, `getenv`, similar).
   - Spawn processes or run shell commands (e.g., `os.system`, `subprocess.*`,
     shell utilities).

2. Perform network or IPC operations:
   - Any HTTP/HTTPS or other network requests.
   - Any sockets, RPC, database, or message-queue operations.

3. Escape or modify the sandbox:
   - Use `eval`, `exec`, `compile`, `execfile`, `globals`, `locals`, `vars` or
     similar dynamic evaluation primitives.
   - Modify `__builtins__` or monkey-patch builtins, modules, or imported
     objects.
   - Use reflection or inspection to access hidden objects (e.g., `inspect`,
     `sys.modules` introspection).

4. Access sensitive or unexpected data:
   - Access absolute or parent directories (e.g., `/etc/`, `/home`, `C:\\`,
     `..` paths).
   - Access secrets or keys, tokens, or configuration files.

5. Abuse control flow or resources:
   - Intentionally infinite or extremely long loops, or obvious attempts to
     consume unbounded CPU/Memory.
   - Very deep recursion or patterns that strongly suggest resource exhaustion
     attacks.

If any of the above are present, the label MUST be `"UNSAFE"`.

## SAFE

Label the code as SAFE when:
- It focuses on creating and formatting resume content inside a `Document`
  using `python-docx` utilities.
- It uses only simple Python control flow (if/else, loops, functions) and basic
  operations on strings, lists, dictionaries, and numbers.
- It optionally imports safe modules such as `docx`, `textwrap`, `math`,
  `statistics`, `itertools`, `collections`.
- It does not attempt any OS, filesystem, environment, network, or sandbox
  escape operations.

## REVIEW

Use the label REVIEW only when the code is unusual or complex, and you cannot
confidently decide SAFE vs UNSAFE given this policy. When in doubt between
SAFE and UNSAFE and the code seems risky, prefer UNSAFE instead of REVIEW.

## EXAMPLES

Example 1 (SAFE):

Content:
"""
from docx import Document

doc = Document()
doc.add_heading("John Doe", level=0)
doc.add_paragraph("Senior Software Engineer")
"""

Expected label: SAFE

Example 2 (UNSAFE - OS access):

Content:
"""
import os
from docx import Document

doc = Document()
secrets = os.listdir("/etc")
doc.add_paragraph(str(secrets))
"""

Expected label: UNSAFE

Example 3 (UNSAFE - subprocess):

Content:
"""
import subprocess
from docx import Document

doc = Document()
output = subprocess.check_output(["ls", "-la"]).decode("utf-8")
doc.add_paragraph(output)
"""

Expected label: UNSAFE
"""


_client: Optional[OpenAI] = None


def _get_client() -> OpenAI:
    """Get or initialize the OpenRouter OpenAI client.

    Raises when the OPENROUTER_API_KEY is missing.
    """
    global _client
    if _client is not None:
        return _client

    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        raise RuntimeError(
            "OPENROUTER_API_KEY is required for DOCX safeguard but is not set."
        )

    _client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    return _client


def _classify_docx_code_sync(code: str) -> Dict[str, Any]:
    """Synchronously classify DOCX Python code according to the policy.

    This is intended to be run in a thread via asyncio.to_thread.
    """
    if not code.strip():
        return {
            "label": "UNSAFE",
            "reasoning": "Empty or whitespace-only code cannot be executed.",
            "violations": ["Empty or whitespace-only code"],
        }

    client = _get_client()

    messages = [
        {
            "role": "system",
            "content": _DOCX_CODE_POLICY,
        },
        {
            "role": "user",
            "content": code,
        },
    ]

    response = client.chat.completions.create(
        model=DOCX_SAFEGUARD_MODEL,
        messages=messages,
        response_format={"type": "json_object"},
    )

    content = response.choices[0].message.content
    if not content:
        raise RuntimeError("DOCX safeguard model returned empty content.")

    try:
        result = json.loads(content)
    except json.JSONDecodeError as exc:  # pragma: no cover - defensive
        raise RuntimeError(
            "DOCX safeguard model returned non-JSON content."
        ) from exc

    return result


async def classify_docx_code_safety(code: str) -> Optional[Dict[str, Any]]:
    """Classify DOCX Python code as SAFE/UNSAFE/REVIEW.

    Returns the parsed model result when the safeguard is enabled. When the
    safeguard is disabled (DOCX_SAFEGUARD_ENABLED=false), returns None so that
    callers can treat the code as unchecked and rely on other defenses.

    Any runtime errors (missing API key, network issues, etc.) are propagated
    to the caller so they can decide whether to fail open or closed.
    """
    if not DOCX_SAFEGUARD_ENABLED:
        return None

    result = await asyncio.to_thread(_classify_docx_code_sync, code)
    return result


def is_label_safe(result: Dict[str, Any]) -> bool:
    """Return True if the safeguard label is SAFE.

    Treat any unexpected or missing labels as non-safe.
    """
    label = str(result.get("label", "")).upper()
    return label == "SAFE"
